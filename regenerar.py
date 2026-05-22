#!/usr/bin/env python3
"""
regenerar.py — Regenera index.html a partir de los .docx de problemas.

Uso:
    cd /Users/nacho/opositor-fq
    python3 regenerar.py

Qué hace:
    1. Lee todos los .docx de la carpeta SECUNDARIA/Problemas
    2. Extrae y limpia los problemas
    3. Convierte la notación a LaTeX
    4. Reescribe index.html con el nuevo banco embebido
"""

import docx, os, json, re, sys
from pathlib import Path

# ─── Rutas ────────────────────────────────────────────────────────────────────
BASE_DOCX = Path("/Users/nacho/Library/Mobile Documents/com~apple~CloudDocs/$MIS_COSAS/SECUNDARIA/Problemas")
OUTPUT    = Path(__file__).parent / "index.html"

# ─── Problemas añadidos manualmente ─────────────────────────────────────────
# Añade aquí nuevos problemas sin necesidad de un .docx
PROBLEMAS_EXTRA = [
    {
        "titulo": "Periodo y velocidad orbital de Marte",
        "bloque": "Gravitación",
        "nivel": "Medio",
        "enunciado": (
            "La Tierra orbita el Sol con $T_T = 1$ año y $r_T = 1$ UA. "
            "Marte tiene $r_M = 1{,}524$ UA. Calcula: "
            "(a) el periodo de Marte en años y en días, "
            "(b) la velocidad orbital de Marte."
        ),
        "resolucion": (
            "(a) 3ª ley de Kepler ($T^2/r^3 = \\mathrm{cte}$): "
            "$T_M^2 = T_T^2 \\cdot (r_M/r_T)^3 = 1^2 \\times (1{,}524)^3 = 3{,}540$ "
            "$\\Rightarrow T_M = \\sqrt{3{,}540} \\approx 1{,}881$ años $= 687$ días "
            "(valor real: 686,97 días). | "
            "(b) $r_M = 1{,}524 \\times 1{,}496\\times10^{11} = 2{,}280\\times10^{11}$ m; "
            "$T_M = 687 \\times 86400 = 5{,}936\\times10^{7}$ s; "
            "$v_M = 2\\pi r_M / T_M \\approx 24\\,130\\ \\mathrm{m/s} \\approx 24{,}1\\ \\mathrm{km/s}$. "
            "Marte va más lento que la Tierra (29,8 km/s) por estar más lejos."
        ),
    },
    # ── Añade más problemas aquí con el mismo formato ──
]

# ─── Extracción de problemas ──────────────────────────────────────────────────

def limpiar(t):
    return re.sub(r'\s+', ' ', t).strip()

def extraer_cuaderno(doc):
    problemas = []
    current = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        m = re.match(r'^(P\d{3})\.\s+(.+)', t)
        if m:
            if current: problemas.append(current)
            current = {"id_str": m.group(1), "titulo": limpiar(m.group(2)),
                       "bloque": "", "nivel": "Medio", "enunciado": "", "resolucion": ""}
        elif current:
            if t.startswith("Bloque:"):
                parts = t.split("|")
                current["bloque"] = parts[0].replace("Bloque:","").strip()
                if len(parts)>1: current["nivel"] = parts[1].replace("Nivel:","").strip()
            elif t.startswith("Enunciado."):
                current["enunciado"] = limpiar(t[len("Enunciado."):])
            elif t.startswith("Estrategia"):
                current["resolucion"] += limpiar(re.sub(r'^Estrategia.*?\.\s*','',t)) + " | "
            elif t.startswith("Resultado"):
                current["resolucion"] += limpiar(re.sub(r'^Resultado.*?\.\s*','',t))
            elif current["enunciado"] and not current["resolucion"]:
                current["enunciado"] += " " + limpiar(t)
    if current: problemas.append(current)
    return problemas

def extraer_tablas(doc, tema):
    problemas = []
    current = None
    for tabla in doc.tables:
        texto = "\n".join(c.text.strip() for row in tabla.rows for c in row.cells if c.text.strip())
        if not texto: continue
        m = re.match(r'^(P\d+)\s*\[(\w+)\]\s*(.*)', texto, re.DOTALL)
        if m:
            if current: problemas.append(current)
            lines = m.group(3).strip().split('\n')
            titulo = lines[0].strip() if lines else ""
            enun = ' '.join(lines[1:]).strip() if len(lines)>1 else titulo
            current = {"id_str": f"{tema}_{m.group(1)}", "titulo": limpiar(titulo),
                       "bloque": tema, "nivel": m.group(2), "enunciado": limpiar(enun), "resolucion": ""}
        elif current and not current["resolucion"]:
            current["resolucion"] = limpiar(texto)
    if current: problemas.append(current)
    return problemas

def extraer_parrafos(doc, tema):
    problemas = []
    current = None
    modo = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        m = re.match(r'^Problema\s+(\d+)\s*[—\-–]\s*(.+)', t)
        if m:
            if current and current.get("enunciado"): problemas.append(current)
            current = {"id_str": f"{tema}_P{m.group(1)}", "titulo": limpiar(m.group(2)),
                       "bloque": tema, "nivel": "Medio", "enunciado": "", "resolucion": ""}
            modo = "enunciado"
        elif current:
            tl = t.lower()
            if any(x in tl for x in ["resolución:", "resolucion:", "paso 1"]):
                modo = "resolucion"
            if modo == "enunciado" and len(current["enunciado"]) < 500:
                if not re.match(r'^(bloque|oposicion|tema|cuerpo|formulario)', tl):
                    current["enunciado"] += " " + limpiar(t)
            elif modo == "resolucion" and len(current["resolucion"]) < 600:
                current["resolucion"] += " " + limpiar(t)
    if current and current.get("enunciado"): problemas.append(current)
    return problemas

BLOQUE_LABELS = {
    "T05-Dinamica":"Dinámica","T06-Rotacion":"Rotación","T07-DinamSist":"Dinámica de sistemas",
    "T08-Gravitacion":"Gravitación","T12-Gases":"Gases ideales","T19-Electrostatica":"Electrostática",
    "T20-Corriente":"Corriente continua","T21-CampoMag":"Campo magnético",
    "T22-Induccion":"Inducción EM","T23-CA":"Corriente alterna","T23b-CA":"Corriente alterna",
}

def get_area(bloque):
    b = bloque.lower()
    if any(x in b for x in ["cinemát","dinámica","gravitac","rotac","sistema"]):
        return "Mecánica"
    if any(x in b for x in ["corriente","electric","magnét","campo","inducción","ondas","óptica"]):
        return "Electromagnetismo"
    if any(x in b for x in ["termod","gas","calor","temperatura"]):
        return "Termodinámica"
    if any(x in b for x in ["quím","ácido","redox","estequio","enlace","cinét","disolu","solub","orgán","termoquím"]):
        return "Química"
    return "Otros"

NIVEL_MAP = {"Basico":"Básico","basico":"Básico","Basic":"Básico","Medio":"Medio","medio":"Medio",
             "Avanzado":"Avanzado","avanzado":"Avanzado","EBAU":"EBAU","ebau":"EBAU",
             "Básico":"Básico","Medium":"Medio","Advanced":"Avanzado"}

CONFIG = {
    "Cuaderno_100_problemas_patron_Fisica_Quimica_Oposiciones.docx": ("cuaderno", "Cuaderno FQ"),
    "Problemas_Tema12_Gases_Ideales.docx": ("tablas",  "T12-Gases"),
    "Problemas_Tema14_Energia_Transferencia.docx": ("parrafos","T14-Energia"),
    "Problemas_Tema20_Corriente_Electrica.docx": ("tablas",  "T20-Corriente"),
    "Problemas_Tema21.docx":  ("parrafos","T21-CampoMag"),
    "Problemas_Tema22.docx":  ("parrafos","T22-Induccion"),
    "Problemas_Tema23_Corriente_Alterna.docx": ("tablas","T23-CA"),
    "Problemas_Tema7_Dinamica_Sistemas.docx":  ("tablas","T07-DinamSist"),
    "Problemas_Tema8_Gravitacion.docx":        ("tablas","T08-Gravitacion"),
    "Problemas_tema14_energia.docx":  ("parrafos","T14b-Energia"),
    "Problemas_tema23.docx":          ("parrafos","T23b-CA"),
    "Problemas_tema5_dinamica.docx":  ("parrafos","T05-Dinamica"),
    "Problemas_tema6_Rotacion.docx":  ("tablas",  "T06-Rotacion"),
    "problemas_Tema19.docx":          ("parrafos","T19-Electrostatica"),
}

def extraer_todos():
    todos = []
    for fname, (modo, tema) in CONFIG.items():
        path = BASE_DOCX / fname
        if not path.exists():
            print(f"  ⚠  No encontrado: {fname}")
            continue
        try:
            d = docx.Document(str(path))
            if modo == "cuaderno":  probs = extraer_cuaderno(d)
            elif modo == "tablas":  probs = extraer_tablas(d, tema)
            else:                   probs = extraer_parrafos(d, tema)
            probs = [p for p in probs if len(p.get("enunciado","")) > 20]
            # Normalizar
            for i, p in enumerate(probs):
                bloque = BLOQUE_LABELS.get(p.get("bloque",""), p.get("bloque",""))
                if not bloque or bloque in ("Cuaderno FQ",): bloque = p.get("bloque","Física y Química")
                nivel = NIVEL_MAP.get(p.get("nivel",""), p.get("nivel","Medio"))
                p.update({"bloque": bloque, "nivel": nivel, "area": get_area(bloque)})
            todos.extend(probs)
            print(f"  ✓ {fname}: {len(probs)} problemas")
        except Exception as e:
            print(f"  ✗ {fname}: {e}")
    # Añadir problemas manuales
    for p in PROBLEMAS_EXTRA:
        extra = dict(p)
        extra.setdefault("nivel", "Medio")
        extra["area"] = get_area(extra.get("bloque", ""))
        todos.append(extra)
    if PROBLEMAS_EXTRA:
        print(f"  ✓ problemas_extra (manual): {len(PROBLEMAS_EXTRA)} problemas")

    # Asignar IDs numéricos únicos
    for i, p in enumerate(todos):
        p["id"] = i
    return todos

# ─── Conversión a LaTeX ───────────────────────────────────────────────────────

GREEK_SUBS = [
    (r'\bomega\b','\\omega'),(r'\btheta\b','\\theta'),(r'\blambda\b','\\lambda'),
    (r'\balpha\b','\\alpha'),(r'\bbeta\b','\\beta'),(r'\bgamma\b','\\gamma'),
    (r'\bdelta\b','\\delta'),(r'\bphi\b','\\phi'),(r'\bsigma\b','\\sigma'),
    (r'\btau\b','\\tau'),(r'\brho\b','\\rho'),(r'\bmu\b','\\mu'),
]
CHEM = [('CO2','CO₂'),('H2O','H₂O'),('O2','O₂'),('N2','N₂'),('NH3','NH₃'),
        ('H2SO4','H₂SO₄'),('H2S','H₂S'),('H2','H₂'),('CH4','CH₄'),('HNO3','HNO₃')]

def M(x): return '$' + x + '$'

def sci_repl(m):
    return M(m.group(1).replace(',','.') + '\\times10^{' + m.group(2) + '}')

def latexify(t, sol=False):
    if not t: return t
    t = re.sub(r'(?<!\d)(\d+[,.]?\d*)[eE]([+-]?\d+)(?!\d)', sci_repl, t)
    for pat, rep in GREEK_SUBS:
        lit = M(rep)
        t = re.sub(pat, lambda m, s=lit: s, t, flags=re.IGNORECASE)
    t = re.sub(r'2\s*pi\b', lambda m: M('2\\pi'), t, flags=re.IGNORECASE)
    t = re.sub(r'([0-9])\s*pi\b', lambda m: M(m.group(1)+'\\pi'), t)
    t = re.sub(r'\bpi\b(?=\s*/)', lambda m: M('\\pi'), t)
    t = re.sub(r'sqrt\(([^)]{1,40})\)', lambda m: M('\\sqrt{'+m.group(1)+'}'), t)
    t = re.sub(r'sqrt(\d+)', lambda m: M('\\sqrt{'+m.group(1)+'}'), t)
    t = re.sub(r'√\(([^)]+)\)', lambda m: M('\\sqrt{'+m.group(1)+'}'), t)
    t = re.sub(r'√(\w+)', lambda m: M('\\sqrt{'+m.group(1)+'}'), t)
    t = re.sub(r'\b([vVsTtRrFfEeBbIiLlMmGgWwZzQqPN])_?([0-9]+)\b',
               lambda m: M(m.group(1)+'_{'+m.group(2)+'}'), t)
    sub_w = r'(?:total|max|min|cm|ef|rms|neto|ext|int|foton|media|reac|fren)'
    t = re.sub(rf'\b([a-zA-Z])_({sub_w})\b',
               lambda m: M(m.group(1)+'_\\mathrm{'+m.group(2)+'}'), t, flags=re.IGNORECASE)
    t = re.sub(r'(?<![/\d])1/2(?![/\d])', lambda m: M('\\tfrac{1}{2}'), t)
    t = re.sub(r'(?<![\\$\{])([a-zA-Z\d])\^([+-]?\d+)(?!\})',
               lambda m: M(m.group(1)+'^{'+m.group(2)+'}'), t)
    t = re.sub(r'\b(min|s|K|mol)\^([+-]\d+)',
               lambda m: M('\\mathrm{'+m.group(1)+'}^{'+m.group(2)+'}'), t)
    for plain, pretty in CHEM:
        t = re.sub(r'\b'+re.escape(plain)+r'\b', pretty, t)
    if sol:
        t = re.sub(
            r'\b([A-Za-z_]{1,4})\s*(≈|=)\s*(-?[\d][,.\d]*)\s*(m/s²?|N|J|W|K|Pa|Hz|rad/s²?|kg·m²|m²?|kΩ|Ω|eV|V|A|C|T)\b',
            lambda m: M(m.group(1)+m.group(2)+m.group(3))+' '+m.group(4), t)
    t = re.sub(r'\$\s*\$', '', t)
    t = re.sub(r'\$([^$]{0,1})\$', r'\1', t)
    return t

# ─── Generación del HTML ──────────────────────────────────────────────────────

def generar_html(banco_js, n_problemas):
    B = chr(96)
    css = """:root{--bg:#0f1117;--bg2:#1a1d2e;--bg3:#232640;--accent:#6c63ff;--accent2:#a78bfa;--orange:#f97316;--green:#22c55e;--yellow:#eab308;--text:#e2e8f0;--text2:#94a3b8;--border:#2d3748;--card:#1e2235;--radius:14px;--shadow:0 4px 24px rgba(0,0,0,.4)}
*{box-sizing:border-box;margin:0;padding:0}
body{background:var(--bg);color:var(--text);font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;min-height:100vh;line-height:1.6}
header{background:linear-gradient(135deg,#1a1d2e,#0f1117);border-bottom:1px solid var(--border);padding:20px 24px 16px;position:sticky;top:0;z-index:100}
.header-top{display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:12px}
.app-title h1{font-size:1.35rem;font-weight:700;background:linear-gradient(135deg,#6c63ff,#a78bfa);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text}
.app-title .sub{font-size:.75rem;color:var(--text2);margin-top:2px}
.stats-row{display:flex;gap:12px;align-items:center;flex-wrap:wrap}
.chip{display:flex;align-items:center;gap:6px;background:var(--bg3);border:1px solid var(--border);border-radius:20px;padding:5px 12px;font-size:.85rem;font-weight:600}
.chip.fire{border-color:#f97316;color:#f97316}.chip.green{border-color:var(--green);color:var(--green)}.chip.purple{border-color:var(--accent2);color:var(--accent2)}
.prog-wrap{margin-top:12px;display:flex;align-items:center;gap:10px}
.prog-bar{flex:1;height:6px;background:var(--bg3);border-radius:3px;overflow:hidden}
.prog-fill{height:100%;background:linear-gradient(90deg,var(--accent),var(--accent2));border-radius:3px;transition:width .5s ease}
.prog-lbl{font-size:.75rem;color:var(--text2);white-space:nowrap}
main{max-width:800px;margin:0 auto;padding:28px 20px 80px}
.motiv{background:linear-gradient(135deg,var(--bg3),var(--bg2));border:1px solid var(--border);border-left:4px solid var(--accent);border-radius:var(--radius);padding:16px 20px;margin-bottom:28px;display:flex;align-items:flex-start;gap:14px}
.motiv-icon{font-size:2rem}.motiv-text h2{font-size:1rem;font-weight:700;color:var(--accent2);margin-bottom:4px}.motiv-text p{font-size:.9rem;color:var(--text2)}
.sec-hdr{display:flex;align-items:center;justify-content:space-between;margin-bottom:18px}
.sec-title{font-size:1.1rem;font-weight:700;display:flex;align-items:center;gap:8px}
.date-badge{font-size:.75rem;color:var(--text2);background:var(--bg3);border:1px solid var(--border);border-radius:8px;padding:3px 8px}
.tabs{display:flex;gap:4px;background:var(--bg2);border:1px solid var(--border);border-radius:10px;padding:4px;margin-bottom:24px}
.tab{flex:1;text-align:center;padding:8px 12px;border-radius:8px;font-size:.85rem;font-weight:600;cursor:pointer;color:var(--text2);transition:all .2s;border:none;background:transparent}
.tab.active{background:var(--accent);color:#fff;box-shadow:0 2px 8px rgba(108,99,255,.4)}
.card{background:var(--card);border:1px solid var(--border);border-radius:var(--radius);margin-bottom:16px;overflow:hidden;transition:border-color .2s,transform .2s}
.card:hover{border-color:var(--accent);transform:translateY(-2px)}
.card.done{border-color:var(--green)!important;opacity:.8}
.card.done .card-hdr{background:rgba(34,197,94,.08)}
.card-hdr{display:flex;align-items:flex-start;justify-content:space-between;padding:14px 18px 10px;gap:12px}
.card-meta{display:flex;align-items:center;gap:8px;flex-wrap:wrap}
.tag-b{font-size:.72rem;font-weight:600;background:rgba(108,99,255,.15);color:var(--accent2);border:1px solid rgba(108,99,255,.3);border-radius:6px;padding:2px 8px;text-transform:uppercase;letter-spacing:.4px}
.tag-n{font-size:.72rem;font-weight:600;border-radius:6px;padding:2px 8px}
.n-Basico,.n-Básico{background:rgba(34,197,94,.15);color:#4ade80;border:1px solid rgba(34,197,94,.3)}
.n-Medio{background:rgba(234,179,8,.15);color:#fbbf24;border:1px solid rgba(234,179,8,.3)}
.n-Avanzado{background:rgba(249,115,22,.15);color:#fb923c;border:1px solid rgba(249,115,22,.3)}
.n-EBAU{background:rgba(239,68,68,.15);color:#f87171;border:1px solid rgba(239,68,68,.3)}
.card-num{font-size:.8rem;color:var(--text2);font-weight:600;white-space:nowrap;margin-top:2px}
.card-body{padding:0 18px 14px}
.card-tit{font-size:1rem;font-weight:700;color:var(--text);margin-bottom:8px}
.card-enun{font-size:.95rem;color:var(--text2);line-height:1.75}
.card-foot{display:flex;gap:10px;align-items:center;padding:10px 18px 14px;border-top:1px solid var(--border);flex-wrap:wrap}
.btn-sol{background:var(--bg3);border:1px solid var(--border);color:var(--text2);border-radius:8px;padding:7px 14px;font-size:.85rem;cursor:pointer;transition:all .2s;display:flex;align-items:center;gap:6px}
.btn-sol:hover{background:var(--accent);border-color:var(--accent);color:#fff}
.btn-ok{background:rgba(34,197,94,.1);border:1px solid rgba(34,197,94,.3);color:#4ade80;border-radius:8px;padding:7px 14px;font-size:.85rem;cursor:pointer;transition:all .2s;display:flex;align-items:center;gap:6px;font-weight:600;margin-left:auto}
.btn-ok:hover{background:rgba(34,197,94,.25)}.btn-ok.done{background:rgba(34,197,94,.2);border-color:var(--green);color:var(--green)}
.sol-panel{display:none;padding:14px 18px;background:rgba(108,99,255,.06);border-top:1px solid rgba(108,99,255,.2);font-size:.9rem;color:var(--text2);line-height:1.8}
.sol-panel.vis{display:block}.sol-panel strong{color:var(--accent2)}
.stats-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(140px,1fr));gap:12px;margin-bottom:28px}
.stat-card{background:var(--card);border:1px solid var(--border);border-radius:var(--radius);padding:16px;text-align:center}
.stat-card .big{font-size:2.2rem;font-weight:800;line-height:1;margin-bottom:4px}
.stat-card .lbl{font-size:.78rem;color:var(--text2);font-weight:500}
.sc-fire .big{color:var(--orange)}.sc-green .big{color:var(--green)}.sc-purple .big{color:var(--accent2)}.sc-yellow .big{color:var(--yellow)}
.cal-grid{display:grid;grid-template-columns:repeat(7,1fr);gap:4px;margin-bottom:12px}
.cal-day{aspect-ratio:1;border-radius:4px;background:var(--bg3)}
.cal-day.prac{background:var(--accent);opacity:.8}.cal-day.today{background:var(--accent2)}
.tema-row{display:flex;align-items:center;gap:10px;margin-bottom:10px}
.tema-lbl{font-size:.82rem;color:var(--text2);width:160px;flex-shrink:0}
.tema-bw{flex:1;height:8px;background:var(--bg3);border-radius:4px;overflow:hidden}
.tema-bf{height:100%;background:linear-gradient(90deg,var(--accent),var(--accent2));border-radius:4px}
.tema-cnt{font-size:.8rem;color:var(--text2);width:40px;text-align:right}
.tab-content{display:none}.tab-content.active{display:block}
.empty{text-align:center;padding:40px 20px;color:var(--text2)}
.empty .icon{font-size:3rem;margin-bottom:12px}
.empty h3{font-size:1.2rem;color:var(--text);margin-bottom:8px}
.toast{position:fixed;bottom:24px;left:50%;transform:translateX(-50%) translateY(80px);background:var(--green);color:#fff;border-radius:12px;padding:12px 20px;font-size:.9rem;font-weight:600;box-shadow:var(--shadow);transition:transform .3s ease;z-index:999}
.toast.show{transform:translateX(-50%) translateY(0)}
.btn-nuevo{background:var(--bg3);border:1px dashed var(--border);color:var(--text2);border-radius:10px;padding:10px 20px;font-size:.85rem;cursor:pointer;transition:all .2s;margin-top:20px}
.btn-nuevo:hover{border-color:var(--accent);color:var(--accent2)}
.filters select{background:var(--bg3);border:1px solid var(--border);color:var(--text);border-radius:8px;padding:6px 10px;font-size:.85rem;cursor:pointer}
.katex{font-size:1em!important;color:var(--text)}
.sol-panel .katex{color:#c4b5fd}
@media(max-width:600px){main{padding:20px 14px 60px}.stats-row{gap:8px}.chip{font-size:.78rem;padding:4px 10px}.card-foot{gap:8px}.btn-ok{margin-left:0}}"""

    js = f"""
const BANCO = {banco_js};
const N_BANCO = {n_problemas};
const ST={{get(k,d){{try{{const v=localStorage.getItem(k);return v!==null?JSON.parse(v):d}}catch(e){{return d}}}},set(k,v){{try{{localStorage.setItem(k,JSON.stringify(v))}}catch(e){{}}}}}};
function cargarEstado(){{return{{resueltos:ST.get('op_resueltos',{{}}),dias:ST.get('op_dias',[]),problemasHoy:ST.get('op_hoy',null)}};}}
let E=cargarEstado();
function hoyStr(){{const d=new Date();return d.getFullYear()+'-'+String(d.getMonth()+1).padStart(2,'0')+'-'+String(d.getDate()).padStart(2,'0');}}
function fechaLegible(s){{const[y,m,d]=s.split('-');const M=['ene','feb','mar','abr','may','jun','jul','ago','sep','oct','nov','dic'];return parseInt(d)+' '+M[parseInt(m)-1]+' '+y;}}
function calcRacha(dias){{if(!dias.length)return 0;const sorted=[...new Set(dias)].sort().reverse();const hoy=hoyStr();const ayer=(()=>{{const d=new Date();d.setDate(d.getDate()-1);return d.toISOString().slice(0,10)}})();if(sorted[0]!==hoy&&sorted[0]!==ayer)return 0;let r=0,esp=sorted[0]===hoy?hoy:ayer;for(const dia of sorted){{if(dia===esp){{r++;const d=new Date(esp);d.setDate(d.getDate()-1);esp=d.toISOString().slice(0,10);}}else break;}}return r;}}
function seededRand(seed){{let s=seed;return function(){{s=(s*1664525+1013904223)&0xffffffff;return(s>>>0)/0xffffffff}};}}
function selProblemasHoy(fecha,forzar){{const saved=E.problemasHoy;if(!forzar&&saved&&saved.fecha===fecha)return saved.ids;const res=E.resueltos;const seed=parseInt(fecha.replace(/-/g,''));const rand=seededRand(seed);let pool=BANCO.filter(p=>!res[p.id]);if(pool.length<5)pool=BANCO;const mezclados=[...pool].sort(()=>rand()-0.5);const areas={{}};const ids=[];for(const p of mezclados){{if(ids.length>=5)break;if((areas[p.area]||0)<2||ids.length>=3){{ids.push(p.id);areas[p.area]=(areas[p.area]||0)+1;}}}}while(ids.length<5&&mezclados.length>ids.length){{const p=mezclados[ids.length];if(!ids.includes(p.id))ids.push(p.id);}}const result={{fecha,ids:ids.slice(0,5)}};ST.set('op_hoy',result);E.problemasHoy=result;return result.ids;}}
function renderKatex(el){{if(typeof renderMathInElement==='undefined')return;try{{renderMathInElement(el,{{delimiters:[{{left:'$',right:'$',display:false}}],throwOnError:false,errorColor:'#f97316'}});}}catch(e){{}}}}
function nClass(n){{return 'n-'+n.replace(/[áéíóú]/g,c=>({{á:'a',é:'e',í:'i',ó:'o',ú:'u'}})[c]||c);}}
function renderCard(p,num,ctx){{const done=!!E.resueltos[p.id];const hasSol=p.resolucion&&p.resolucion.length>5;return {B}<div class="card${{done?' done':''}}" id="card-${{ctx}}-${{p.id}}"><div class="card-hdr"><div class="card-meta"><span class="tag-b">${{p.bloque}}</span><span class="tag-n ${{nClass(p.nivel)}}">${{p.nivel}}</span></div><span class="card-num">#${{String(num).padStart(3,'0')}}</span></div><div class="card-body">${{p.titulo?{B}<div class="card-tit">${{p.titulo}}</div>{B}:''}}<div class="card-enun">${{p.enunciado}}</div></div><div class="card-foot">${{hasSol?{B}<button class="btn-sol" onclick="toggleSol('${{ctx}}',${{p.id}})"><span id="si-${{ctx}}-${{p.id}}">👁</span> Ver solución</button>{B}:''}}<button class="btn-ok${{done?' done':''}}" onclick="marcar(${{p.id}},'${{ctx}}',this)">${{done?'✅ Resuelto':'○ Marcar resuelto'}}</button></div>${{hasSol?{B}<div class="sol-panel" id="sol-${{ctx}}-${{p.id}}"><strong>Solución:</strong> ${{p.resolucion}}</div>{B}:''}}</div>{B};}}
function toggleSol(ctx,id){{const panel=document.getElementById(`sol-${{ctx}}-${{id}}`);const icon=document.getElementById(`si-${{ctx}}-${{id}}`);if(panel){{panel.classList.toggle('vis');if(icon)icon.textContent=panel.classList.contains('vis')?'🙈':'👁';if(panel.classList.contains('vis'))renderKatex(panel);}}}}
function marcar(id,ctx,btn){{const res=ST.get('op_resueltos',{{}});const hoy=hoyStr();if(res[id]){{delete res[id];btn.classList.remove('done');btn.textContent='○ Marcar resuelto';const c=document.getElementById(`card-${{ctx}}-${{id}}`);if(c)c.classList.remove('done');}}else{{res[id]=hoy;btn.classList.add('done');btn.innerHTML='✅ Resuelto';const c=document.getElementById(`card-${{ctx}}-${{id}}`);if(c)c.classList.add('done');const dias=ST.get('op_dias',[]);if(!dias.includes(hoy)){{dias.push(hoy);ST.set('op_dias',dias);E.dias=dias;}}toast('¡Problema resuelto! 🎉');}}ST.set('op_resueltos',res);E.resueltos=res;updateHeader();}}
function updateHeader(){{const res=E.resueltos;const dias=ST.get('op_dias',[]);const racha=calcRacha(dias);const total=Object.keys(res).length;const ids=E.problemasHoy?.ids||[];const hoyN=ids.filter(id=>res[id]).length;document.getElementById('v-racha').textContent=racha;document.getElementById('v-total').textContent=total;document.getElementById('v-hoy').textContent=hoyN;document.getElementById('prog-fill').style.width=(hoyN/5*100)+'%';document.getElementById('prog-lbl').textContent=hoyN+' de 5 problemas de hoy';updateMotiv(racha,hoyN,total);const rm=ST.get('op_rachaMax',0);if(racha>rm)ST.set('op_rachaMax',racha);}}
function updateMotiv(racha,hoyN,total){{const msgs=[{{c:()=>hoyN===5,i:'🏆',t:'¡Día completado!',p:'Los 5 problemas de hoy dominados. ¡Imparable!'}},{{c:()=>racha>=7,i:'🔥',t:`¡${{racha}} días seguidos!`,p:'Una racha brutal. El tribunal no sabe lo que le viene.'}},{{c:()=>hoyN>=3,i:'💪',t:'¡Más de la mitad, a por el resto!',p:'Ya has hecho lo difícil. Completa el día.'}},{{c:()=>racha>=3,i:'⚡',t:`Racha de ${{racha}} días`,p:'La constancia es la clave del éxito en oposiciones.'}},{{c:()=>total>=50,i:'🧠',t:`${{total}} problemas resueltos`,p:'Ya tienes una base sólida. Sigue sumando.'}},{{c:()=>true,i:'🚀',t:'¡A por el día!',p:'Tienes 5 problemas esperándote. Cada uno te acerca al tribunal.'}}];const m=msgs.find(x=>x.c());document.getElementById('motiv-i').textContent=m.i;document.getElementById('motiv-t').textContent=m.t;document.getElementById('motiv-p').textContent=m.p;}}
function renderHoy(){{const hoy=hoyStr();document.getElementById('fecha-badge').textContent=fechaLegible(hoy);const ids=selProblemasHoy(hoy);const cont=document.getElementById('lista-hoy');cont.innerHTML=ids.map(id=>{{const p=BANCO.find(x=>x.id===id);return p?renderCard(p,id+1,'hoy'):''}}).join('');renderKatex(cont);updateHeader();}}
function nuevosDia(){{if(!confirm('¿Generar 5 nuevos problemas aleatorios?'))return;const hoy=hoyStr();const res=E.resueltos;let pool=BANCO.filter(p=>!res[p.id]);if(pool.length<5)pool=BANCO;const ids=[...pool].sort(()=>Math.random()-0.5).slice(0,5).map(p=>p.id);const saved={{fecha:hoy,ids}};ST.set('op_hoy',saved);E.problemasHoy=saved;renderHoy();}}
function filtrarBanco(){{const area=document.getElementById('f-area').value,nivel=document.getElementById('f-nivel').value,ef=document.getElementById('f-estado').value,res=E.resueltos;let lista=BANCO;if(area)lista=lista.filter(p=>p.area===area);if(nivel)lista=lista.filter(p=>p.nivel===nivel);if(ef==='pendiente')lista=lista.filter(p=>!res[p.id]);if(ef==='resuelto')lista=lista.filter(p=>!!res[p.id]);document.getElementById('banco-cnt').textContent=lista.length+' problemas';const cont=document.getElementById('lista-banco');if(!lista.length){{cont.innerHTML='<div class="empty"><div class="icon">🎉</div><h3>¡Todos resueltos!</h3><p>Has completado todos los problemas de este filtro.</p></div>';return;}}cont.innerHTML=lista.slice(0,60).map(p=>renderCard(p,p.id+1,'banco')).join('');renderKatex(cont);}}
function renderProgreso(){{const res=E.resueltos,dias=ST.get('op_dias',[]),racha=calcRacha(dias),rachaMax=Math.max(racha,ST.get('op_rachaMax',0));document.getElementById('s-racha').textContent=racha;document.getElementById('s-rmax').textContent=rachaMax;document.getElementById('s-total').textContent=Object.keys(res).length;document.getElementById('s-banco').textContent=N_BANCO;const cal=document.getElementById('calendario');cal.innerHTML='';const diasSet=new Set(dias);const hoy=new Date();for(let i=27;i>=0;i--){{const d=new Date(hoy);d.setDate(d.getDate()-i);const s=d.toISOString().slice(0,10);const div=document.createElement('div');div.className='cal-day'+(i===0?' today':(diasSet.has(s)?' prac':''));div.title=s;cal.appendChild(div);}}const tm={{}};for(const p of BANCO){{if(!tm[p.bloque])tm[p.bloque]={{total:0,r:0}};tm[p.bloque].total++;if(res[p.id])tm[p.bloque].r++;}}const sorted=Object.entries(tm).sort((a,b)=>b[1].total-a[1].total);document.getElementById('temas-prog').innerHTML=sorted.map(([t,{{total,r}}])=>{B}<div class="tema-row"><div class="tema-lbl">${{t}}</div><div class="tema-bw"><div class="tema-bf" style="width:${{Math.round(r/total*100)}}%"></div></div><div class="tema-cnt">${{r}}/${{total}}</div></div>{B}).join('');}}
function mostrarTab(tab,el){{document.querySelectorAll('.tab-content').forEach(x=>x.classList.remove('active'));document.querySelectorAll('.tab').forEach(x=>x.classList.remove('active'));document.getElementById('tab-'+tab).classList.add('active');el.classList.add('active');if(tab==='banco')filtrarBanco();if(tab==='progreso')renderProgreso();}}
function toast(msg){{const t=document.getElementById('toast-el');t.textContent=msg;t.classList.add('show');setTimeout(()=>t.classList.remove('show'),2500);}}
function resetProgress(){{if(!confirm('¿Borrar todo el progreso?'))return;['op_resueltos','op_dias','op_rachaMax','op_hoy'].forEach(k=>localStorage.removeItem(k));E=cargarEstado();renderHoy();toast('Progreso borrado');}}
document.addEventListener('DOMContentLoaded',()=>renderHoy());"""

    return f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<meta http-equiv="Cache-Control" content="no-cache, no-store, must-revalidate">
<title>Opositor Diario — Física y Química</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.css">
<style>{css}</style>
</head>
<body>
<header>
  <div class="header-top">
    <div class="app-title"><h1>⚛️ Opositor Diario</h1><div class="sub">Física y Química · Cuerpo 590 · CyL</div></div>
    <div class="stats-row">
      <div class="chip fire">🔥 <span id="v-racha">0</span> días</div>
      <div class="chip green">✓ <span id="v-total">0</span> resueltos</div>
      <div class="chip purple">⭐ <span id="v-hoy">0</span>/5 hoy</div>
    </div>
  </div>
  <div class="prog-wrap"><div class="prog-bar"><div class="prog-fill" id="prog-fill" style="width:0%"></div></div><div class="prog-lbl" id="prog-lbl">0 de 5 problemas de hoy</div></div>
</header>
<main>
  <div class="tabs">
    <button class="tab active" onclick="mostrarTab('hoy',this)">📅 Hoy</button>
    <button class="tab" onclick="mostrarTab('banco',this)">📚 Banco</button>
    <button class="tab" onclick="mostrarTab('progreso',this)">📊 Progreso</button>
  </div>
  <div class="tab-content active" id="tab-hoy">
    <div class="motiv"><div class="motiv-icon" id="motiv-i">🚀</div><div class="motiv-text"><h2 id="motiv-t">¡A por el día!</h2><p id="motiv-p">Tienes 5 problemas esperándote. Cada uno te acerca al tribunal.</p></div></div>
    <div class="sec-hdr"><div class="sec-title">📝 Problemas de hoy</div><div class="date-badge" id="fecha-badge">—</div></div>
    <div id="lista-hoy"></div>
    <div style="text-align:center"><button class="btn-nuevo" onclick="nuevosDia()">↻ Nuevos 5 problemas</button></div>
  </div>
  <div class="tab-content" id="tab-banco">
    <div class="sec-hdr"><div class="sec-title">📚 Todos los problemas</div><div class="date-badge" id="banco-cnt">— problemas</div></div>
    <div class="filters" style="display:flex;gap:8px;margin-bottom:16px;flex-wrap:wrap;">
      <select id="f-area" onchange="filtrarBanco()"><option value="">Todas las áreas</option><option value="Mecánica">Mecánica</option><option value="Electromagnetismo">Electromagnetismo</option><option value="Termodinámica">Termodinámica</option><option value="Química">Química</option></select>
      <select id="f-nivel" onchange="filtrarBanco()"><option value="">Todos los niveles</option><option value="Básico">Básico</option><option value="Medio">Medio</option><option value="Avanzado">Avanzado</option><option value="EBAU">EBAU</option></select>
      <select id="f-estado" onchange="filtrarBanco()"><option value="">Todos</option><option value="pendiente">Pendientes</option><option value="resuelto">Resueltos</option></select>
    </div>
    <div id="lista-banco"></div>
  </div>
  <div class="tab-content" id="tab-progreso">
    <div class="stats-grid">
      <div class="stat-card sc-fire"><div class="big" id="s-racha">0</div><div class="lbl">🔥 Racha actual</div></div>
      <div class="stat-card sc-fire"><div class="big" id="s-rmax">0</div><div class="lbl">🏆 Racha máxima</div></div>
      <div class="stat-card sc-green"><div class="big" id="s-total">0</div><div class="lbl">✓ Total resueltos</div></div>
      <div class="stat-card sc-yellow"><div class="big" id="s-banco">{n_problemas}</div><div class="lbl">📚 En banco</div></div>
    </div>
    <div class="sec-title" style="margin-bottom:14px">📅 Últimas 4 semanas</div>
    <div class="cal-grid" id="calendario"></div>
    <div style="font-size:.75rem;color:var(--text2);margin-bottom:24px">
      <span style="display:inline-block;width:12px;height:12px;background:var(--accent2);border-radius:2px;margin-right:4px;vertical-align:middle"></span>Hoy
      <span style="display:inline-block;width:12px;height:12px;background:var(--accent);border-radius:2px;margin-right:4px;margin-left:12px;vertical-align:middle"></span>Practicaste
      <span style="display:inline-block;width:12px;height:12px;background:var(--bg3);border-radius:2px;margin-right:4px;margin-left:12px;vertical-align:middle"></span>Sin práctica
    </div>
    <div class="sec-title" style="margin-bottom:14px">🎯 Por tema</div>
    <div id="temas-prog"></div>
    <div style="margin-top:28px"><div class="sec-title" style="margin-bottom:12px">🗑️ Reiniciar</div>
      <button onclick="resetProgress()" style="background:rgba(239,68,68,.1);border:1px solid rgba(239,68,68,.3);color:#f87171;border-radius:8px;padding:8px 16px;font-size:.85rem;cursor:pointer">Borrar todo el progreso</button>
    </div>
  </div>
</main>
<div class="toast" id="toast-el"></div>
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.js"></script>
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/contrib/auto-render.min.js"></script>
<script>
{js}
</script>
</body>
</html>"""

# ─── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=" * 60)
    print("Opositor Diario — Regenerando banco de problemas")
    print("=" * 60)

    # 1. Verificar dependencia
    try:
        import docx
    except ImportError:
        print("Instalando python-docx...")
        os.system("pip3 install python-docx -q")
        import docx

    # 2. Extraer
    print("\n📂 Extrayendo problemas de los .docx...")
    problemas = extraer_todos()
    print(f"\n  → {len(problemas)} problemas extraídos en total")

    # 3. Convertir a LaTeX (solo los extraídos de docx, no los manuales)
    ids_extra = {id(p) for p in PROBLEMAS_EXTRA}
    print("\n🔢 Convirtiendo notación a LaTeX...")
    for p in problemas:
        if id(p) not in ids_extra:          # los manuales ya vienen en LaTeX
            p["enunciado"]  = latexify(p.get("enunciado",""),  sol=False)
            p["resolucion"] = latexify(p.get("resolucion",""), sol=True)
        # Limpiar campos internos
        for k in ["id_str"]:
            p.pop(k, None)

    # 4. Generar HTML
    print("\n🔨 Generando index.html...")
    banco_js = json.dumps(problemas, ensure_ascii=False, separators=(',',':'))
    html = generar_html(banco_js, len(problemas))

    OUTPUT.write_text(html, encoding="utf-8")
    print(f"  → Escrito: {len(html):,} bytes en {OUTPUT}")

    # 5. Git
    print("\n📤 Publicando en GitHub...")
    os.chdir(OUTPUT.parent)
    os.system('git add index.html')
    os.system(f'git commit -m "Actualizar banco: {len(problemas)} problemas"')
    os.system('git push')

    print(f"""
✅ ¡Listo!
   • {len(problemas)} problemas en el banco
   • Publicado en https://jilucasf.github.io/opositor-fq/
     (disponible en 1-2 minutos)
""")
